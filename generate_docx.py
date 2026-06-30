import collections
import collections.abc
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_heading(doc, text, level=1, color=RGBColor(94, 53, 177)):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return heading

def add_bullet(doc, text, bold_prefix="", indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.name = 'Calibri'
        run_b.font.color.rgb = RGBColor(33, 33, 33)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(68, 68, 68)
    return p

def add_body_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(50, 50, 50)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Add a horizontal line via bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_pitch_doc():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ==========================================
    # COVER PAGE
    # ==========================================
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('C2B PathFinder')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(94, 53, 177)
    run.font.name = 'Calibri'

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_p.add_run('AI-Powered Learning Path Recommendation Platform')
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(0, 137, 123)
    run2.font.name = 'Calibri'

    doc.add_paragraph()

    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = desc_p.add_run(
        'Bridging the gap between Campus and Business with explainable, '
        'data-driven skill mapping and intelligent course recommendations.'
    )
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(100, 100, 100)
    run3.font.name = 'Calibri'

    for _ in range(4):
        doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = footer_p.add_run('BGT 4th Edition  •  Pitch Presentation  •  2026')
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(150, 150, 150)
    run4.font.name = 'Calibri'

    doc.add_page_break()

    # ==========================================
    # SLIDE 2: THE TEAM
    # ==========================================
    add_styled_heading(doc, '1. The Team', level=1)
    add_body_text(doc, 
        'Our multidisciplinary team combines AI engineering, full-stack development, '
        'and product design expertise to deliver a cohesive, intelligent platform.'
    )

    # Team table
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    headers = ['Lead AI Engineer', 'Full-Stack Developer', 'Product & UX Architect']
    subtitles = ['Algorithms & Models', 'Platform Engineering', 'Strategic Design']
    skills = [
        '• Machine Learning & NLP\n• Recommendation Architectures\n• Explainable AI (XAI)\n• Python & PyTorch',
        '• React 18, Vite & Node.js\n• Tailwind CSS & Dashboards\n• API Integration\n• CI/CD & Deployment',
        '• Interactive Heatmap Design\n• Feedback Loop Workflows\n• User Journey Modeling\n• B2B Analytics Patterns'
    ]

    for i in range(3):
        cell = table.cell(0, i)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(headers[i])
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(94, 53, 177)
        run.font.name = 'Calibri'
        set_cell_shading(cell, 'F3F0FF')

    for i in range(3):
        cell = table.cell(1, i)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(subtitles[i])
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 137, 123)
        run.font.name = 'Calibri'

    for i in range(3):
        cell = table.cell(2, i)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(skills[i])
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(68, 68, 68)

    # Merge row 3 for a note
    merged = table.cell(3, 0).merge(table.cell(3, 2))
    merged.text = ''
    p = merged.paragraphs[0]
    run = p.add_run('Each team member brings domain-specific expertise that directly maps to a pillar of the platform.')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.font.name = 'Calibri'

    doc.add_page_break()

    # ==========================================
    # SLIDE 3: THE OPPORTUNITY
    # ==========================================
    add_styled_heading(doc, '2. The Opportunity', level=1)

    add_styled_heading(doc, 'The Problem', level=2, color=RGBColor(94, 53, 177))
    add_bullet(doc, ' Manual skill gap assessments are static, costly, and quickly outdated.', '•')
    add_bullet(doc, ' Traditional training paths are generic and fail to align with individual learner goals.', '•')
    add_bullet(doc, ' Companies lack real-time visibility into department-level capability gaps, leading to wasted training budgets.', '•')
    add_bullet(doc, ' 77% of companies report difficulty finding employees with the right digital skills (World Economic Forum, 2025).', '•')

    doc.add_paragraph()

    add_styled_heading(doc, 'The Opportunity', level=2, color=RGBColor(0, 137, 123))
    add_bullet(doc, ' Leverage AI to dynamically map historical learning paths and predict the "Next Best Course."', '•')
    add_bullet(doc, ' Transform individual Learning Passports into organizational heatmap diagnostics.', '•')
    add_bullet(doc, ' Close the loop using feedback engines to continuously optimize recommendations.', '•')
    add_bullet(doc, ' The global corporate e-learning market is projected to reach $50B by 2026, growing at 15% CAGR.', '•')

    doc.add_page_break()

    # ==========================================
    # SLIDE 4: THE SOLUTION
    # ==========================================
    add_styled_heading(doc, '3. The Solution: C2B PathFinder', level=1)
    add_body_text(doc, 
        'C2B PathFinder integrates three conceptual pillars into one unified platform with three dashboards:'
    )

    # Pillar 1
    add_styled_heading(doc, 'Pillar 1: Learning Passport (Learner Dashboard)', level=2, color=RGBColor(94, 53, 177))
    add_body_text(doc, 
        'A dynamic, living record tracking completed courses, acquired skills, and progress toward a '
        'target career role. Features include:'
    )
    add_bullet(doc, ' Visual skill radar chart showing competency across multiple dimensions')
    add_bullet(doc, ' AI-generated course recommendations with confidence scores and plain-language explanations')
    add_bullet(doc, ' Interactive Learning Path visualization from current role to target role')
    add_bullet(doc, ' Gap analysis showing acquired vs. remaining target skills')

    doc.add_paragraph()

    # Pillar 2
    add_styled_heading(doc, 'Pillar 2: Skills Gap Diagnostic (Company Dashboard)', level=2, color=RGBColor(0, 137, 123))
    add_body_text(doc, 
        'Translates individual learner data into a department-wide skills heatmap for enterprise administrators:'
    )
    add_bullet(doc, ' Color-coded grid showing departments vs. skills with gap counts (e.g., "20/25 employees lack this skill")')
    add_bullet(doc, ' Critical gap alerts highlighting urgent training needs')
    add_bullet(doc, ' AI-generated phased training roadmap aligned with budget (e.g., €120,000) and timeline (12 months)')
    add_bullet(doc, ' Real-time progress tracking across all departments')

    doc.add_paragraph()

    # Pillar 3
    add_styled_heading(doc, 'Pillar 3: Feedback Loop Engine (Admin Dashboard)', level=2, color=RGBColor(126, 87, 194))
    add_body_text(doc, 
        'Monitors post-course outcomes and continuously optimizes the recommendation model:'
    )
    add_bullet(doc, ' Post-course feedback with impact scores and real-world application rates')
    add_bullet(doc, ' Platform-wide analytics: 1,247 learners, 38 companies, 30 courses, 78% average completion')
    add_bullet(doc, ' Course catalogue management with search, filtering, and metadata')
    add_bullet(doc, ' Dropout risk alerts and completion trend monitoring')

    doc.add_page_break()

    # ==========================================
    # SLIDE 5: MARKET FIT & DIFFERENTIATION
    # ==========================================
    add_styled_heading(doc, '4. Market Fit & Differentiation', level=1)

    add_styled_heading(doc, 'Core Differentiators', level=2, color=RGBColor(94, 53, 177))
    add_bullet(doc, ' Unlike black-box systems, learners see explicit confidence metrics and reasons '
                     '(e.g., "recommended because it fills 3 remaining skill gaps toward your target role").', 
                'Explainable AI (XAI): ')
    add_bullet(doc, ' The platform is self-correcting. Post-course feedback automatically refines '
                     'and improves the recommendation engine over time.',
                'Closed-Loop Learning: ')
    add_bullet(doc, ' Scalable dashboards that bridge individual employee upskilling with '
                     'company-wide strategic workforce planning seamlessly.',
                'Multi-Scale Architecture: ')

    doc.add_paragraph()

    add_styled_heading(doc, 'Market Alignment', level=2, color=RGBColor(0, 137, 123))
    add_bullet(doc, ' Connects vocational training institutes directly with their commercial enterprise partners, '
                     'creating a two-sided marketplace.',
                'Academic/B2B Synergy: ')
    add_bullet(doc, ' Providing explicit, transparent career growth roadmaps boosts course completion rates by over 30%.',
                'Employee Retention: ')
    add_bullet(doc, ' Empowers corporate L&D administrators to spend budgets only on critical gaps '
                     'identified by the department heatmap.',
                'L&D Budget Efficiency: ')

    doc.add_page_break()

    # ==========================================
    # SLIDE 6: BUSINESS MODEL
    # ==========================================
    add_styled_heading(doc, '5. Business Model', level=1)
    add_body_text(doc, 'C2B PathFinder is designed for long-term sustainability through three revenue streams:')

    # Business model table
    biz_table = doc.add_table(rows=4, cols=3)
    biz_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    biz_table.style = 'Light Grid Accent 1'

    biz_headers = ['B2B SaaS License', 'Catalogue Commission', 'Enterprise Add-ons']
    biz_subtitles = ['Recurring Subscription', 'Transaction Fees', 'Custom Integrations']
    biz_descs = [
        'Charged per employee/seat for enterprise and SME clients. Grants full access to company dashboards, '
        'department heatmaps, automated roadmaps, and analytics reports.',
        'Revenue-share model with educational providers and training centers integrated in the catalogue. '
        'Earn commissions for every course enrollment booked through PathFinder.',
        'High-margin service for custom integration with existing enterprise HR/LMS software. '
        'Private fine-tuning of the AI model on corporate-specific competency frameworks.'
    ]

    for i in range(3):
        cell = biz_table.cell(0, i)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(biz_headers[i])
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(94, 53, 177)
        run.font.name = 'Calibri'
        set_cell_shading(cell, 'F3F0FF')

    for i in range(3):
        cell = biz_table.cell(1, i)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(biz_subtitles[i])
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 137, 123)
        run.font.name = 'Calibri'

    for i in range(3):
        cell = biz_table.cell(2, i)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(biz_descs[i])
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(68, 68, 68)

    merged = biz_table.cell(3, 0).merge(biz_table.cell(3, 2))
    merged.text = ''
    p = merged.paragraphs[0]
    run = p.add_run('Combined, these streams ensure recurring revenue while scaling with customer growth.')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.font.name = 'Calibri'

    doc.add_page_break()

    # ==========================================
    # SLIDE 7: FINAL REQUEST
    # ==========================================
    add_styled_heading(doc, '6. The Request', level=1)

    for _ in range(2):
        doc.add_paragraph()

    ask_p = doc.add_paragraph()
    ask_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ask_p.add_run('Partner with C2B PathFinder')
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(94, 53, 177)
    run.font.name = 'Calibri'

    ask_sub = doc.add_paragraph()
    ask_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = ask_sub.add_run("Let's build the future of dynamic skills diagnostics together.")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0, 137, 123)
    run2.font.name = 'Calibri'

    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    add_styled_heading(doc, 'Our Request:', level=2, color=RGBColor(0, 137, 123))

    add_bullet(doc, ' We are seeking Corporate & Training Pilot Partners to deploy the C2B PathFinder '
                     'platform with active learner cohorts and validate real-world impact.', '1.')
    add_bullet(doc, ' We are looking for advisory partnerships with L&D leaders to optimize our '
                     "model's taxonomy matching and competency framework alignment.", '2.')
    add_bullet(doc, ' Explore our live platform and source code:', '3.')

    doc.add_paragraph()

    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_link = link_p.add_run('Live Demo: https://annkahoro.github.io/C2B_PROJECT/')
    run_link.font.size = Pt(11)
    run_link.bold = True
    run_link.font.color.rgb = RGBColor(94, 53, 177)
    run_link.font.name = 'Calibri'

    link_p2 = doc.add_paragraph()
    link_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_link2 = link_p2.add_run('Repository: https://github.com/Annkahoro/C2B_PROJECT')
    run_link2.font.size = Pt(11)
    run_link2.bold = True
    run_link2.font.color.rgb = RGBColor(0, 137, 123)
    run_link2.font.name = 'Calibri'

    for _ in range(4):
        doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer.add_run('C2B PathFinder — BGT 4th Edition — Campus2Business Challenge — 2026')
    run_f.italic = True
    run_f.font.size = Pt(10)
    run_f.font.color.rgb = RGBColor(150, 150, 150)
    run_f.font.name = 'Calibri'

    # Save
    doc.save('C2B_PathFinder_Pitch.docx')
    print('Word document saved to C2B_PathFinder_Pitch.docx')

if __name__ == '__main__':
    create_pitch_doc()
